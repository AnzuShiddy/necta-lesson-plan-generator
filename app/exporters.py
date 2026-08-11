"""Render a LessonPlanDocument to .docx and .pdf byte buffers."""

import io

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from .schema import LessonPlanDocument


# Two lesson plan layouts share one page skeleton:
#
#   "classic"  — the school-file layout: LESSON PLAN title, School/Teacher,
#                Form/Subject, Time/Date lines, then the common body.
#   "tie2023"  — a facsimile of the pre-printed TIE lesson plan booklet: no
#                title and no school/teacher lines, Subject/Class then Date/Time
#                paired down the page, Main Competence, the Registered/Present
#                students table, then Specific Competence, Main Activity,
#                Specific Activities, Resources and References.
#
# Both then print the "Teaching and Learning Process" table and Remarks.
_LP_DOTS = "." * 30
_PROCESS_COLS = ["Stages", "Time (Minutes)", "Teaching Activities",
                 "Learning Activities", "Assessment Criteria"]
# Column proportions follow the printed TIE booklet page, except that Assessment
# Criteria is widened from the booklet's 12% to 16%. The booklet is written on by
# hand in a few words; a generated criterion runs to ~100 characters, and at the
# printed width it wrapped to ten lines and drove the whole table onto a second
# page. Time only ever holds a number, so it gives up the space.
_PROCESS_WIDTHS = [3.0, 1.8, 5.4, 5.4, 3.0]  # cm, sums to the A4 text width


def _is_tie(doc: LessonPlanDocument) -> bool:
    return doc.request.plan_format == "tie2023"


def _lp_time(doc: LessonPlanDocument) -> str:
    r = doc.request
    time = r.time
    if r.period_number:
        time = f"Period {r.period_number}, {time}".strip(", ")
    if r.duration_minutes:
        time = f"{time} ({r.duration_minutes} minutes)".strip()
    return time


def _lp_header_lines(doc: LessonPlanDocument) -> list[str]:
    r = doc.request
    return [
        f"Name of School: {r.school_name or _LP_DOTS}"
        f"        Teacher's Name: {r.teacher_name or _LP_DOTS}",
        f"Form: {f'{r.form} {r.stream}'.strip()}        Subject: {r.subject}",
        f"Time: {_lp_time(doc) or _LP_DOTS}        Date: {r.date or _LP_DOTS}",
    ]


def _tie_header_pairs(doc: LessonPlanDocument) -> list[tuple[str, str]]:
    """The four boxed-in header fields of the TIE booklet, in printed order."""
    r = doc.request
    return [
        ("Subject", r.subject),
        ("Class", f"{r.form} {r.stream}".strip() or _LP_DOTS),
        ("Date", r.date or _LP_DOTS),
        ("Time", _lp_time(doc) or _LP_DOTS),
    ]


def _lp_fields(doc: LessonPlanDocument) -> list[tuple[str, str]]:
    """Body fields printed below the students table.

    On the TIE form Main Competence sits *above* the students table, so it is
    emitted separately there; the classic layout keeps it here at the top.
    """
    p = doc.plan
    resources = "; ".join(p.teaching_learning_resources)
    references = "; ".join(p.references)
    main_activity = doc.request.subtopic or p.lesson_title
    specific = "; ".join(p.specific_activities) or p.lesson_title

    if _is_tie(doc):
        return [
            ("Specific Competence", p.specific_competence),
            ("Main Activity", main_activity),
            ("Specific Activities", specific),
            ("Teaching and Learning Resources", resources),
            ("References", references),
        ]

    fields = [
        ("Main Competence", p.main_competence),
        ("Specific Competence", p.specific_competence),
        ("Main Activity", main_activity),
        ("Specific Activities", specific),
    ]
    if doc.request.week_label:
        fields.append(("Scheme of Work Reference", doc.request.week_label))
    fields += [
        ("Teaching and Learning Resources", resources),
        ("References", references),
    ]
    return fields


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def _docx_students_table(d: Document, doc: LessonPlanDocument):
    """Number of students: Registered vs Present (present left for the teacher)."""
    r = doc.request
    stable = d.add_table(rows=4, cols=6)
    stable.style = "Table Grid"
    stable.alignment = WD_TABLE_ALIGNMENT.CENTER
    top = stable.rows[0].cells
    head = top[0].merge(top[5]).paragraphs[0]
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    head.add_run("Number of students").bold = True
    mid = stable.rows[1].cells
    for cells, label in ((mid[0].merge(mid[2]), "Registered"),
                         (mid[3].merge(mid[5]), "Present")):
        para = cells.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run(label).bold = True
    for c, label in zip(stable.rows[2].cells, ["Girls", "Boys", "Total"] * 2):
        para = c.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run(label).bold = True
    for c, value in zip(stable.rows[3].cells,
                        [r.girls, r.boys, r.boys + r.girls, "", "", ""]):
        c.text = str(value)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return stable


def _docx_process_table(d: Document, doc: LessonPlanDocument):
    table = d.add_table(rows=1, cols=len(_PROCESS_COLS))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for cell, label, width in zip(table.rows[0].cells, _PROCESS_COLS,
                                  _PROCESS_WIDTHS):
        cell.width = Cm(width)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run(label).bold = True
    for st in doc.plan.stages:
        cells = table.add_row().cells
        for cell, width, text in zip(
            cells, _PROCESS_WIDTHS,
            [st.stage, str(st.duration_minutes), st.teaching_activities,
             st.learning_activities, st.assessment],
        ):
            cell.width = Cm(width)
            cell.text = text
        cells[0].paragraphs[0].runs[0].bold = True
    return table


def to_docx(doc: LessonPlanDocument) -> bytes:
    d = Document()
    style = d.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    if _is_tie(doc):
        for sec in d.sections:
            sec.left_margin = sec.right_margin = Cm(1.2)

    def kv(label: str, value: str):
        p = d.add_paragraph()
        p.add_run(label + ": ").bold = True
        p.add_run(value)
        return p

    if _is_tie(doc):
        # The booklet page carries no title and no school/teacher lines: it
        # opens straight into Subject/Class, Date/Time, then Main Competence.
        pairs = _tie_header_pairs(doc)
        head = d.add_table(rows=2, cols=2)   # borderless, to align the columns
        head.autofit = False
        for row, (left, right) in zip(head.rows, [pairs[:2], pairs[2:]]):
            for cell, (label, value) in zip(row.cells, (left, right)):
                cell.width = Cm(9.3)
                para = cell.paragraphs[0]
                para.add_run(label + ": ").bold = True
                para.add_run(value)
        kv("Main Competence", doc.plan.main_competence)
        d.add_paragraph()
        _docx_students_table(d, doc)
        d.add_paragraph()
        for label, value in _lp_fields(doc):
            kv(label, value)
        d.add_paragraph()
        heading = d.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.add_run("Teaching and Learning Process").bold = True
    else:
        title = d.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("LESSON PLAN")
        run.bold = True
        run.font.size = Pt(14)
        for line in _lp_header_lines(doc):
            d.add_paragraph(line)
        _docx_students_table(d, doc)
        d.add_paragraph()
        for label, value in _lp_fields(doc):
            kv(label, value)
        d.add_paragraph()
        d.add_paragraph().add_run("Teaching and Learning Process").bold = True

    _docx_process_table(d, doc)

    d.add_paragraph()
    kv("Remarks", _LP_DOTS)   # written by hand after the lesson — see schema.py

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def _pdf_students_table(doc: LessonPlanDocument, style: ParagraphStyle,
                        tie: bool) -> Table:
    r = doc.request
    total = r.boys + r.girls
    centred = ParagraphStyle("stucell", parent=style, alignment=1)
    data = [
        [Paragraph("<b>Number of students</b>", centred), "", "", "", "", ""],
        [Paragraph("<b>Registered</b>", centred), "", "",
         Paragraph("<b>Present</b>", centred), "", ""],
        [Paragraph(f"<b>{h}</b>", centred) for h in ["Girls", "Boys", "Total"] * 2],
        [Paragraph(str(v), centred) for v in [r.girls, r.boys, total]] + ["", "", ""],
    ]
    width = 2.6 * cm if tie else 2.2 * cm
    table = Table(data, colWidths=[width] * 6,
                  hAlign="CENTER" if tie else "LEFT")
    table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
        ("SPAN", (0, 0), (5, 0)),
        ("SPAN", (0, 1), (2, 1)),
        ("SPAN", (3, 1), (5, 1)),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 2),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
    ]))
    return table


def _pdf_process_table(doc: LessonPlanDocument, cell: ParagraphStyle,
                       head: ParagraphStyle, tie: bool) -> Table:
    data = [[Paragraph(h, head) for h in _PROCESS_COLS]]
    for st in doc.plan.stages:
        data.append([
            Paragraph(f"<b>{st.stage}</b>", cell),
            Paragraph(str(st.duration_minutes), cell),
            Paragraph(st.teaching_activities.replace("\n", "<br/>"), cell),
            Paragraph(st.learning_activities.replace("\n", "<br/>"), cell),
            Paragraph(st.assessment.replace("\n", "<br/>"), cell),
        ])
    table = Table(data, colWidths=[w * cm for w in _PROCESS_WIDTHS], repeatRows=1)
    style = [
        ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]
    # The printed booklet has no shaded header row; the classic layout keeps
    # the app's green banner.
    if not tie:
        style.append(("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f7a4d")))
    table.setStyle(TableStyle(style))
    return table


def to_pdf(doc: LessonPlanDocument) -> bytes:
    tie = _is_tie(doc)
    buf = io.BytesIO()
    pdf = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=1.2 * cm,
        rightMargin=1.2 * cm,
        topMargin=1.2 * cm,
        bottomMargin=1.2 * cm,
        title="Lesson Plan",
    )
    styles = getSampleStyleSheet()
    cell = ParagraphStyle("cell", parent=styles["Normal"], fontSize=8, leading=10)
    cell_head = ParagraphStyle(
        "cellhead", parent=cell, fontName="Helvetica-Bold", alignment=1,
        textColor=colors.black if tie else colors.white)
    h1 = ParagraphStyle("h1", parent=styles["Title"], fontSize=16)
    small = ParagraphStyle("small", parent=styles["Normal"], fontSize=9, leading=14)
    centre = ParagraphStyle("centre", parent=small, alignment=1, fontSize=11,
                            fontName="Helvetica-Bold")

    story: list = []

    def block(label, value):
        story.append(Paragraph(f"<b>{label}:</b> {value}", small))

    if tie:
        # Subject/Class and Date/Time paired across the page, as printed.
        pairs = _tie_header_pairs(doc)
        rows = [[Paragraph(f"<b>{lbl}:</b> {val}", small) for lbl, val in pair]
                for pair in (pairs[:2], pairs[2:])]
        header = Table(rows, colWidths=[9.3 * cm, 9.3 * cm], hAlign="LEFT")
        header.setStyle(TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 0),
            ("TOPPADDING", (0, 0), (-1, -1), 0),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
        ]))
        story.append(header)
        block("Main Competence", doc.plan.main_competence)
        story.append(Spacer(1, 8))
        story.append(_pdf_students_table(doc, small, tie))
        story.append(Spacer(1, 10))
        for label, value in _lp_fields(doc):
            block(label, value)
        story.append(Spacer(1, 8))
        story.append(Paragraph("Teaching and Learning Process", centre))
    else:
        story.append(Paragraph("LESSON PLAN", h1))
        for line in _lp_header_lines(doc):
            story.append(Paragraph(line.replace("        ", "&nbsp;" * 8), small))
        story.append(Spacer(1, 6))
        story.append(_pdf_students_table(doc, small, tie))
        story.append(Spacer(1, 8))
        for label, value in _lp_fields(doc):
            block(label, value)
        story.append(Spacer(1, 8))
        story.append(Paragraph("<b>Teaching and Learning Process</b>", small))
    story.append(Spacer(1, 4))

    story.append(_pdf_process_table(doc, cell, cell_head, tie))
    story.append(Spacer(1, 8))
    block("Remarks", _LP_DOTS)   # written by hand after the lesson — see schema.py

    pdf.build(story)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Scheme of work exporters (the scheme is itself a document teachers submit)
# ---------------------------------------------------------------------------

# Official TIE (2023 revised curriculum) scheme-of-work layout: landscape page,
# SCHEME OF WORK title, School/Teacher/Subject/Year/Term header lines, then a
# 12-column table, one table per term.
_SCHEME_COLS = [
    ("Main Competence", 2.6),
    ("Specific Competences", 2.8),
    ("Learning Activities", 3.4),
    ("Specific Activities", 2.0),
    ("Month", 1.4),
    ("Week", 0.9),
    ("Number of Periods", 1.0),
    ("Teaching and Learning Methods", 4.1),
    ("Teaching and Learning Resources", 2.6),
    ("Assessment Tools", 2.4),
    ("References", 2.8),
    ("Remarks", 1.4),
]
_TERM_NAMES = {1: "TERM I", 2: "TERM II"}
_DOTS = "." * 44


def _scheme_row(e: dict) -> list[str]:
    # Specific Activities is the teacher's own breakdown of the learning
    # activity; pre-fill only the multi-week split so the teacher completes it.
    specific = ""
    if e.get("topic_weeks", 1) > 1:
        specific = f"Part {e['topic_week']} of {e['topic_weeks']} of the learning activity"
    return [
        e.get("main_competence", ""),
        e.get("specific_competence", ""),
        e.get("learning_activity", ""),
        specific,
        e["month"],
        str(e["week"]),
        str(e["periods"]),
        "; ".join(e.get("teaching_learning_activities", [])),
        ", ".join(e.get("resources", [])),
        e.get("assessment", ""),
        e.get("references", ""),
        e.get("remarks", ""),
    ]


def _scheme_terms(sch: dict) -> list[tuple[str, list[dict]]]:
    terms: list[tuple[str, list[dict]]] = []
    for sem in (1, 2):
        entries = [e for e in sch["entries"] if e["semester"] == sem]
        if entries:
            terms.append((_TERM_NAMES[sem], entries))
    return terms


def scheme_to_docx(sch: dict) -> bytes:
    d = Document()
    sec = d.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    sec.left_margin = sec.right_margin = Cm(1)
    d.styles["Normal"].font.name = "Calibri"
    d.styles["Normal"].font.size = Pt(7)

    title = d.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SCHEME OF WORK")
    run.bold = True
    run.font.size = Pt(14)
    for line in (
        f"Name of School: {_DOTS}        Teacher's Name: {_DOTS}",
        f"Subject: {sch['subject']}        Form: {sch['form']}",
        f"Year: {sch['year']}        Periods per week: {sch['periods_per_week']}",
    ):
        p = d.add_paragraph()
        p.add_run(line).font.size = Pt(10)

    for term_name, entries in _scheme_terms(sch):
        hp = d.add_paragraph()
        hr = hp.add_run(term_name)
        hr.bold = True
        hr.font.size = Pt(11)
        table = d.add_table(rows=1, cols=len(_SCHEME_COLS))
        table.style = "Table Grid"
        table.autofit = False
        for c, (label, width) in zip(table.rows[0].cells, _SCHEME_COLS):
            c.width = Cm(width)
            c.paragraphs[0].add_run(label).bold = True
        for e in entries:
            cells = table.add_row().cells
            for c, (_, width), text in zip(cells, _SCHEME_COLS, _scheme_row(e)):
                c.width = Cm(width)
                c.text = text

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def scheme_to_pdf(sch: dict) -> bytes:
    buf = io.BytesIO()
    pdf = SimpleDocTemplate(buf, pagesize=landscape(A4), leftMargin=1 * cm,
                            rightMargin=1 * cm, topMargin=1 * cm, bottomMargin=1 * cm,
                            title="Scheme of Work")
    styles = getSampleStyleSheet()
    head = ParagraphStyle("shead", parent=styles["Normal"], fontSize=10, leading=14)
    cell = ParagraphStyle("scell", parent=styles["Normal"], fontSize=6.5, leading=7.5)
    cell_b = ParagraphStyle("scellb", parent=cell, fontName="Helvetica-Bold",
                            textColor=colors.white)
    gap = "&nbsp;" * 8
    story = [
        Paragraph("SCHEME OF WORK", styles["Title"]),
        Paragraph(f"Name of School: {_DOTS}{gap}Teacher's Name: {_DOTS}", head),
        Paragraph(f"Subject: {sch['subject']}{gap}Form: {sch['form']}", head),
        Paragraph(f"Year: {sch['year']}{gap}Periods per week: {sch['periods_per_week']}", head),
        Spacer(1, 6),
    ]
    for term_name, entries in _scheme_terms(sch):
        story.append(Paragraph(term_name, styles["Heading3"]))
        data = [[Paragraph(label, cell_b) for label, _ in _SCHEME_COLS]]
        for e in entries:
            data.append([Paragraph(text, cell) for text in _scheme_row(e)])
        table = Table(data, colWidths=[w * cm for _, w in _SCHEME_COLS], repeatRows=1)
        table.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f7a4d")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("TOPPADDING", (0, 0), (-1, -1), 2),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
        ]))
        story.append(table)
        story.append(Spacer(1, 10))
    pdf.build(story)
    return buf.getvalue()
